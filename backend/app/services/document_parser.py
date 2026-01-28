"""
Document parsing service for content extraction.

Handles different file types and extracts text content for indexing.
Uses Docling for PDFs/DOCX, ElevenLabs for audio transcription.
"""
import logging
import tempfile
import os
from typing import List, Optional, Dict, Any, Tuple, TYPE_CHECKING
from pathlib import Path

from app.core.config import settings
from app.models.document import FileType

if TYPE_CHECKING:
    from app.services.elevenlabs_service import ElevenLabsService
    from app.services.docling_service import DoclingService

logger = logging.getLogger(__name__)

# Chunking configuration
DEFAULT_CHUNK_SIZE = 1000  # characters
DEFAULT_CHUNK_OVERLAP = 200  # characters
MAX_CHUNK_SIZE = 2000


class DocumentParserService:
    """Service for parsing documents and extracting text content."""

    def __init__(
        self,
        elevenlabs_service: Optional["ElevenLabsService"] = None,
        docling_service: Optional["DoclingService"] = None,
    ):
        """
        Initialize the parser service.

        Args:
            elevenlabs_service: Optional ElevenLabs service instance.
                               If not provided, will be created lazily when needed.
            docling_service: Optional Docling service instance.
                            If not provided, will be created lazily when needed.
        """
        self._elevenlabs_service = elevenlabs_service
        self._docling_service = docling_service

    @property
    def elevenlabs_service(self) -> "ElevenLabsService":
        """Get or create ElevenLabs service instance."""
        if self._elevenlabs_service is None:
            from app.services.elevenlabs_service import get_elevenlabs_service
            self._elevenlabs_service = get_elevenlabs_service()
        return self._elevenlabs_service

    @property
    def docling_service(self) -> "DoclingService":
        """Get or create Docling service instance."""
        if self._docling_service is None:
            from app.services.docling_service import get_docling_service
            self._docling_service = get_docling_service()
        return self._docling_service

    async def parse_document(
        self,
        file_data: bytes,
        file_type: str,
        filename: str,
    ) -> Tuple[str, Dict[str, Any]]:
        """
        Parse document and extract text content.

        Args:
            file_data: Raw file bytes
            file_type: FileType enum value
            filename: Original filename

        Returns:
            Tuple of (extracted_text, metadata_dict)
        """
        file_type_enum = FileType(file_type)

        if file_type_enum == FileType.PDF:
            return await self._parse_pdf(file_data, filename)
        elif file_type_enum in (FileType.DOCX, FileType.DOC):
            return await self._parse_docx(file_data, filename)
        elif file_type_enum in (FileType.TXT, FileType.MD):
            return await self._parse_text(file_data, filename)
        elif file_type_enum in (FileType.AUDIO_MP3, FileType.AUDIO_WAV, FileType.AUDIO_M4A):
            return await self._parse_audio(file_data, filename, file_type_enum)
        elif file_type_enum in (FileType.VIDEO_MP4, FileType.VIDEO_WEBM):
            return await self._parse_video(file_data, filename, file_type_enum)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")

    async def _parse_pdf(
        self,
        file_data: bytes,
        filename: str,
    ) -> Tuple[str, Dict[str, Any]]:
        """Parse PDF using Docling service."""
        try:
            result = self.docling_service.parse_bytes(file_data, filename)

            text = result.get("text", "")

            metadata = {
                "page_count": result.get("page_count"),
                "table_count": result.get("table_count", 0),
                "parser": "docling",
                "has_markdown": result.get("markdown") is not None,
            }

            return text, metadata

        except ImportError:
            logger.warning("Docling not installed, falling back to PyPDF2")
            return await self._parse_pdf_fallback(file_data, filename)
        except Exception as e:
            logger.warning(f"Docling parsing failed: {e}, falling back to PyPDF2")
            return await self._parse_pdf_fallback(file_data, filename)

    async def _parse_pdf_fallback(
        self,
        file_data: bytes,
        filename: str,
    ) -> Tuple[str, Dict[str, Any]]:
        """Fallback PDF parsing using PyPDF2."""
        try:
            from PyPDF2 import PdfReader
            import io

            reader = PdfReader(io.BytesIO(file_data))
            text_parts = []

            for page in reader.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)

            text = "\n\n".join(text_parts)
            metadata = {
                "page_count": len(reader.pages),
                "parser": "pypdf2",
            }

            return text, metadata

        except ImportError:
            raise ValueError("No PDF parser available. Install docling or PyPDF2.")

    async def _parse_docx(
        self,
        file_data: bytes,
        filename: str,
    ) -> Tuple[str, Dict[str, Any]]:
        """Parse DOCX/DOC files using Docling service."""
        try:
            result = self.docling_service.parse_bytes(file_data, filename)

            text = result.get("text", "")

            metadata = {
                "page_count": result.get("page_count"),
                "table_count": result.get("table_count", 0),
                "parser": "docling",
                "has_markdown": result.get("markdown") is not None,
            }

            return text, metadata

        except ImportError:
            logger.warning("Docling not installed, falling back to python-docx")
            return await self._parse_docx_fallback(file_data, filename)
        except Exception as e:
            logger.warning(f"Docling DOCX parsing failed: {e}, falling back to python-docx")
            return await self._parse_docx_fallback(file_data, filename)

    async def _parse_docx_fallback(
        self,
        file_data: bytes,
        filename: str,
    ) -> Tuple[str, Dict[str, Any]]:
        """Fallback DOCX parsing using python-docx."""
        try:
            from docx import Document
            import io

            doc = Document(io.BytesIO(file_data))
            text_parts = []

            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        text_parts.append(row_text)

            text = "\n\n".join(text_parts)
            metadata = {
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(doc.tables),
                "parser": "python-docx",
            }

            return text, metadata

        except ImportError:
            raise ValueError("python-docx not installed. Install it to parse DOCX files.")

    async def _parse_text(
        self,
        file_data: bytes,
        filename: str,
    ) -> Tuple[str, Dict[str, Any]]:
        """Parse plain text files."""
        # Try different encodings
        for encoding in ["utf-8", "latin-1", "cp1252"]:
            try:
                text = file_data.decode(encoding)
                metadata = {
                    "encoding": encoding,
                    "line_count": text.count("\n") + 1,
                    "parser": "text",
                }
                return text, metadata
            except UnicodeDecodeError:
                continue

        raise ValueError("Failed to decode text file with supported encodings")

    async def _parse_audio(
        self,
        file_data: bytes,
        filename: str,
        file_type: FileType,
    ) -> Tuple[str, Dict[str, Any]]:
        """Transcribe audio using ElevenLabs SDK."""
        if not settings.ELEVENLABS_API_KEY:
            raise ValueError("ELEVENLABS_API_KEY required for audio transcription")

        try:
            result = await self.elevenlabs_service.transcribe_audio(
                audio_data=file_data,
                filename=filename,
                model_id="scribe_v1",
            )

            text = result.get("text", "")

            metadata = {
                "parser": "elevenlabs",
                "model": result.get("model", "scribe_v1"),
                "language": result.get("language"),
                "duration": result.get("duration"),
            }

            # Add segments count if available
            if "segments" in result:
                metadata["segment_count"] = len(result["segments"])

            return text, metadata

        except Exception as e:
            logger.error(f"ElevenLabs transcription error: {e}")
            raise ValueError(f"Audio transcription failed: {e}")

    async def _parse_video(
        self,
        file_data: bytes,
        filename: str,
        file_type: FileType,
    ) -> Tuple[str, Dict[str, Any]]:
        """Extract and transcribe audio from video."""
        # For video, we extract audio track and transcribe
        # This requires ffmpeg for audio extraction
        try:
            import subprocess

            # Write video to temp file
            suffix = ".mp4" if file_type == FileType.VIDEO_MP4 else ".webm"
            with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as tmp_video:
                tmp_video.write(file_data)
                video_path = tmp_video.name

            # Extract audio to temp file
            audio_path = video_path + ".mp3"

            try:
                # Use ffmpeg to extract audio
                result = subprocess.run(
                    [
                        "ffmpeg",
                        "-i", video_path,
                        "-vn",  # No video
                        "-acodec", "libmp3lame",
                        "-q:a", "2",  # Quality
                        "-y",  # Overwrite
                        audio_path,
                    ],
                    capture_output=True,
                    timeout=300,  # 5 minute timeout
                )

                if result.returncode != 0:
                    logger.error(f"FFmpeg error: {result.stderr.decode()}")
                    raise ValueError("Failed to extract audio from video")

                # Read extracted audio
                with open(audio_path, "rb") as f:
                    audio_data = f.read()

                # Transcribe audio
                text, audio_metadata = await self._parse_audio(
                    audio_data,
                    filename + ".mp3",
                    FileType.AUDIO_MP3,
                )

                metadata = {
                    **audio_metadata,
                    "source_type": "video",
                    "video_format": file_type.value,
                }

                return text, metadata

            finally:
                os.unlink(video_path)
                if os.path.exists(audio_path):
                    os.unlink(audio_path)

        except FileNotFoundError:
            raise ValueError("FFmpeg not installed. Required for video processing.")

    def chunk_text(
        self,
        text: str,
        chunk_size: int = DEFAULT_CHUNK_SIZE,
        chunk_overlap: int = DEFAULT_CHUNK_OVERLAP,
    ) -> List[Dict[str, Any]]:
        """
        Split text into overlapping chunks.

        Args:
            text: Text to chunk
            chunk_size: Maximum characters per chunk
            chunk_overlap: Characters to overlap between chunks

        Returns:
            List of dicts with 'text' and 'chunk_index'
        """
        if not text or not text.strip():
            return []

        # Clean text
        text = text.strip()

        # If text is short enough, return as single chunk
        if len(text) <= chunk_size:
            return [{"text": text, "chunk_index": 0}]

        chunks = []
        start = 0
        chunk_index = 0

        while start < len(text):
            end = start + chunk_size

            # If not at end, try to break at sentence/paragraph boundary
            if end < len(text):
                # Look for paragraph break
                para_break = text.rfind("\n\n", start, end)
                if para_break > start + chunk_size // 2:
                    end = para_break

                # Otherwise look for sentence break
                elif "." in text[start:end]:
                    sentence_break = text.rfind(". ", start, end)
                    if sentence_break > start + chunk_size // 2:
                        end = sentence_break + 1

                # Otherwise look for any newline
                elif "\n" in text[start:end]:
                    line_break = text.rfind("\n", start, end)
                    if line_break > start + chunk_size // 2:
                        end = line_break

            chunk_text = text[start:end].strip()
            if chunk_text:
                chunks.append({
                    "text": chunk_text,
                    "chunk_index": chunk_index,
                })
                chunk_index += 1

            # Move start with overlap
            start = end - chunk_overlap
            if start >= len(text):
                break

        return chunks


# Singleton instance
_parser_service: Optional[DocumentParserService] = None


def get_document_parser() -> DocumentParserService:
    """Get or create document parser service singleton."""
    global _parser_service
    if _parser_service is None:
        _parser_service = DocumentParserService()
    return _parser_service
